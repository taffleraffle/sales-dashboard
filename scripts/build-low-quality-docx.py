"""Builds Low-Quality-Replacement-List.docx from preview-audit.csv.

Produces a clean Word document with:
- Cover summary
- BROKEN_PLACEHOLDER table (top priority)
- SUB_PAR table
- SQL query block for pulling live URLs
- Notes on caveats
"""

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent
CSV_PATH = ROOT / "preview-audit.csv"
OUT_PATH = ROOT / "Low-Quality-Replacement-List.docx"

URL_PATTERN = (
    "https://kjfaqhmllagbxjdxlopm.supabase.co/storage/v1/object/public/"
    "creative-uploads/previews/<filename>"
)

SQL_QUERY = """SELECT id,
       COALESCE(canonical_name, name) AS name,
       type,
       preview_url,
       preview_url || '?download=' || COALESCE(canonical_name, name) AS download_url,
       low_quality_reason,
       low_quality_actual_mb
FROM lib_creative_library
WHERE is_low_quality = true
ORDER BY low_quality_reason, name;"""


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, font_size=10, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_mono_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def build_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr_cells[i], h, bold=True, font_size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr_cells[i], "1F4E79")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r):
            set_cell_text(row_cells[i], str(val), font_size=9)
        if col_widths:
            for i, w in enumerate(col_widths):
                row_cells[i].width = w
    return table


def main():
    if not CSV_PATH.exists():
        raise SystemExit(f"Missing CSV at {CSV_PATH}")

    with CSV_PATH.open("r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        all_rows = list(reader)

    broken = [r for r in all_rows if r["classification"] == "BROKEN_PLACEHOLDER"]
    subpar = [r for r in all_rows if r["classification"] == "SUB_PAR"]
    ok = [r for r in all_rows if r["classification"] == "OK"]

    doc = Document()

    # Set default font
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    title = doc.add_heading("Low-Quality Creative Library Replacement List", level=0)
    for run in title.runs:
        run.font.name = "Calibri"

    add_paragraph(doc, "Sales Dashboard — Editor Dashboard creative library", italic=True, size=10)
    add_paragraph(doc, f"Source: preview-audit.csv  ·  Project: kjfaqhmllagbxjdxlopm", italic=True, size=10)

    # Summary
    add_heading(doc, "Summary", level=1)
    add_paragraph(
        doc,
        f"{len(broken) + len(subpar)} of {len(all_rows)} audited clips need replacement.",
        bold=True,
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"BROKEN_PLACEHOLDER (actual < 3 MB): ").bold = True
    p.add_run(f"{len(broken)} rows. Truncated download, unplayable.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"SUB_PAR (bitrate < 4 Mbps): ").bold = True
    p.add_run(f"{len(subpar)} rows. Plays but WhatsApp-call quality.")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"OK: ").bold = True
    p.add_run(f"{len(ok)} rows (excluded from this list).")

    # Caveats
    add_heading(doc, "Caveats", level=2)
    add_paragraph(
        doc,
        "1. This is a snapshot from when the audit was last run. Rows that have already "
        "been replaced may still appear. Re-run scripts/audit-preview-file-sizes.mjs for "
        "current state.",
        size=10,
    )
    add_paragraph(
        doc,
        "2. The CSV does not store the URL. Predicted URL pattern is below — the actual "
        "preview_url is in lib_creative_library. Use the SQL query at the end of this doc "
        "to pull live URLs from Supabase.",
        size=10,
    )

    # URL pattern
    add_heading(doc, "Predicted URL pattern", level=2)
    add_mono_block(doc, URL_PATTERN)
    add_paragraph(
        doc,
        "Append ?download=<filename> to force a real binary download "
        "(per the Video Quality Contract).",
        italic=True,
        size=9,
    )

    # BROKEN
    doc.add_page_break()
    add_heading(doc, f"BROKEN_PLACEHOLDER ({len(broken)} rows — top priority)", level=1)
    add_paragraph(
        doc,
        "Actual file < 3 MB. Completely unplayable. Original ingest truncated the download.",
        italic=True,
        size=10,
    )

    broken_rows = []
    for i, r in enumerate(broken, start=1):
        broken_rows.append([
            i,
            r["id"],
            r["name"].strip('"'),
            r["type"],
            f"{float(r['actual_mb']):.2f}",
            "yes" if r["has_drive"] == "true" else "no",
        ])
    build_table(
        doc,
        ["#", "ID", "Name", "Type", "Actual MB", "Has Drive"],
        broken_rows,
        col_widths=[Cm(0.7), Cm(4.2), Cm(6.5), Cm(1.8), Cm(1.4), Cm(1.6)],
    )

    drive_count = sum(1 for r in broken if r["has_drive"] == "true")
    if drive_count:
        add_paragraph(
            doc,
            f"{drive_count} of these have drive_url populated — the original may be "
            "re-fetchable from Drive without manual upload.",
            italic=True,
            size=10,
        )

    # SUB_PAR
    doc.add_page_break()
    add_heading(doc, f"SUB_PAR ({len(subpar)} rows)", level=1)
    add_paragraph(
        doc,
        "Plays but bitrate < 4 Mbps (WhatsApp-call quality).",
        italic=True,
        size=10,
    )

    subpar_rows = []
    for i, r in enumerate(subpar, start=1):
        mbps = r["mbps"].strip() if r["mbps"] else "—"
        subpar_rows.append([
            i,
            r["id"],
            r["name"].strip('"'),
            r["type"],
            mbps,
            f"{float(r['actual_mb']):.2f}",
            "yes" if r["has_drive"] == "true" else "no",
        ])
    build_table(
        doc,
        ["#", "ID", "Name", "Type", "Mbps", "Actual MB", "Has Drive"],
        subpar_rows,
        col_widths=[Cm(0.7), Cm(4.2), Cm(6.0), Cm(1.8), Cm(1.2), Cm(1.4), Cm(1.4)],
    )

    # SQL query
    doc.add_page_break()
    add_heading(doc, "Pull live URLs from Supabase", level=1)
    add_paragraph(
        doc,
        "Paste this into the Supabase SQL editor for project kjfaqhmllagbxjdxlopm "
        "to get the actual preview_url and a ready-to-share download URL for every "
        "flagged row.",
        size=10,
    )
    add_mono_block(doc, SQL_QUERY)

    # Next steps
    add_heading(doc, "Recommended next steps", level=1)
    steps = [
        ("Confirm current state", "Re-run node scripts/audit-preview-file-sizes.mjs (needs SUPABASE_SERVICE_ROLE_KEY in env)."),
        ("Pull live URLs", "Run the SQL above. The download_url column is the link to share."),
        ("Replace from local source", "node scripts/replace-from-local-files.mjs (TUS resumable upload)."),
        ("Check Drive first", "Rows with has_drive = yes may be re-fetchable without re-uploading."),
    ]
    for label, desc in steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(label + ": ")
        run.bold = True
        p.add_run(desc)

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    print(f"  BROKEN_PLACEHOLDER: {len(broken)}")
    print(f"  SUB_PAR:            {len(subpar)}")
    print(f"  OK (excluded):      {len(ok)}")


if __name__ == "__main__":
    main()
